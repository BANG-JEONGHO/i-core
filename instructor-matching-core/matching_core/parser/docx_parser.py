"""DOCX 문서 파서."""

from __future__ import annotations

import io

import structlog

from matching_core.exceptions import EmptyDocumentError, ParseError
from matching_core.models.entities import FileType, ParsedDocument
from matching_core.parser.base import BaseParser
from matching_core.utils.text_processing import clean_text

logger = structlog.get_logger()


class DocxParser(BaseParser):
    """DOCX(Word) 파일을 파싱합니다."""

    def parse(self, file_content: bytes) -> ParsedDocument:
        """DOCX에서 텍스트를 추출합니다."""
        try:
            from docx import Document
        except ImportError as e:
            raise ParseError(
                user_message="DOCX 파싱 라이브러리가 설치되지 않았습니다.",
                internal_detail=f"python-docx import failed: {e}",
            )

        try:
            doc = Document(io.BytesIO(file_content))
            paragraphs: list[str] = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # 테이블 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            raw_text = "\n".join(paragraphs)
            raw_text = clean_text(raw_text)

            if not raw_text.strip():
                raise EmptyDocumentError()

            logger.info(
                "docx_parsed",
                paragraphs=len(paragraphs),
                text_length=len(raw_text),
            )

            return ParsedDocument(
                raw_text=raw_text,
                file_type=FileType.DOCX,
                sections=[],
            )

        except EmptyDocumentError:
            raise
        except Exception as e:
            logger.error("docx_parse_failed", error=str(e))
            raise ParseError(
                user_message="DOCX 파일을 파싱할 수 없습니다. 파일이 손상되었을 수 있습니다.",
                internal_detail=f"DOCX parse error: {e}",
            )
