"""AI Agent - Gemini 기반 과업지시서 분석 및 매칭 추천.

파싱 전략:
- DOCX: python-docx로 텍스트 추출 → Gemini에 텍스트 전달
- HWP: pyhwpx로 텍스트 추출 시도 → 실패 시 Gemini에 원문 전달
- PDF: Gemini File API에 직접 PDF 업로드 → AI가 파싱
"""

from __future__ import annotations

import json
import tempfile
import os
import re
import structlog
from google import genai
from google.genai import types

from app.core.config import settings

logger = structlog.get_logger()

_client: genai.Client | None = None
MODEL = settings.GEMINI_MODEL


def _get_client() -> genai.Client:
    global _client
    if _client is None:
        _client = genai.Client(
            api_key=settings.GEMINI_API_KEY,
            http_options=types.HttpOptions(
                client_args={"trust_env": settings.GEMINI_USE_ENV_PROXY}
            ),
        )
    return _client


EXTRACTION_PROMPT = """당신은 나라장터 교육 과업지시서를 분석하는 전문가입니다.
아래 과업지시서에서 '신청자격(참여자격)', '평가기준', 그리고 매칭에 사용할 '과업 개요'를 추출해주세요.

## 규칙:
1. 신청자격: 참여하기 위한 자격 요건 (학력, 경력, 자격증, 인력 요건 등)
2. 평가기준: 제안서 평가 항목 (기술평가, 가격평가, 실적 평가 등 + 배점)
3. 각 항목에서 핵심 키워드를 추출
4. 필수/우대 구분
5. 과업 개요는 '개요', '사업 개요', '과업 개요', '사업 목적', '추진 배경 및 목적',
   '사업 배경', '과업 범위', '교육 개요', '사업 내용', '추진 내용' 등 문서마다 다른 제목을
   같은 개요 맥락으로 판단합니다. 목차가 아니라 실제 본문만 사용합니다.
6. 개요에는 문서에 명시된 사실만 담고, 추정하지 않습니다. 없는 값은 빈 배열 또는 빈 문자열로 둡니다.
7. source_excerpt에는 개요 판단의 근거가 되는 원문 문장을 짧게 그대로 넣습니다.

## 출력 형식 (JSON만 출력):
{
  "qualifications": [
    {"category": "자격증", "description": "정보처리기사 이상", "is_mandatory": true, "keywords": ["정보처리기사"]}
  ],
  "evaluation_criteria": [
    {"category": "기술평가", "description": "유사 사업 수행 실적", "weight": 60, "keywords": ["유사사업", "실적"]}
  ],
  "overview": {
    "section_titles": ["사업 개요", "사업 목적"],
    "summary": "AWS 및 생성형 AI 실습형 교육을 운영하는 과업",
    "core_topics": ["AWS", "생성형 AI"],
    "target_audience": ["취업준비생"],
    "delivery_formats": ["오프라인", "실습형"],
    "required_roles": ["강의", "프로젝트 멘토링"],
    "scope": ["교육과정 운영", "교안 개발"],
    "expected_outcomes": ["실무 역량 향상"],
    "preferred_experience": ["유사 교육과정 운영 경험"],
    "source_excerpt": "원문에 실제로 있는 개요 근거 문장"
  }
}

JSON만 출력하세요:"""


OVERVIEW_SECTION_ALIASES = (
    "개요",
    "사업개요",
    "과업개요",
    "교육개요",
    "사업목적",
    "과업목적",
    "추진배경및목적",
    "추진배경",
    "사업배경",
    "과업범위",
    "사업내용",
    "추진내용",
)

_OVERVIEW_BOUNDARY_TERMS = (
    "개요", "목적", "배경", "범위", "내용", "추진", "일정", "자격", "평가",
    "제안", "과업", "사업", "요구사항", "수행", "교육", "계약", "예산",
)


def _normalize_heading(value: str) -> str:
    value = value.strip()
    value = re.sub(
        r"^\s*(?:(?:제?\s*\d+\s*[장절.]?)|(?:[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+)|(?:\d+(?:\.\d+)*[.)]?)|(?:[가-하][.)]))\s*",
        "",
        value,
    )
    return re.sub(r"[^0-9A-Za-z가-힣]", "", value).casefold()


def _is_heading_line(value: str) -> bool:
    stripped = value.strip()
    if not stripped or len(stripped) > 70 or stripped.endswith((".", "다.", "요.")):
        return False
    normalized = _normalize_heading(stripped)
    return (
        bool(re.match(r"^\s*(?:[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+|\d+(?:\.\d+)*|[가-하])(?:[.)]|\s)", stripped))
        or any(term in normalized for term in _OVERVIEW_BOUNDARY_TERMS)
    )


def extract_overview_context(raw_text: str) -> dict:
    """Find overview sections by heading aliases without requiring another LLM call."""
    empty = {
        "section_titles": [],
        "summary": "",
        "core_topics": [],
        "target_audience": [],
        "delivery_formats": [],
        "required_roles": [],
        "scope": [],
        "expected_outcomes": [],
        "preferred_experience": [],
        "source_excerpt": "",
        "matching_rules": [],
    }
    if not raw_text or not raw_text.strip():
        return empty

    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    sections: list[tuple[str, str]] = []
    for index, line in enumerate(lines):
        heading = _normalize_heading(line)
        if not heading or not any(heading == alias or heading.startswith(alias) for alias in OVERVIEW_SECTION_ALIASES):
            continue

        content_lines: list[str] = []
        for following_line in lines[index + 1:]:
            if content_lines and _is_heading_line(following_line):
                break
            content_lines.append(following_line)
            if len("\n".join(content_lines)) >= 2_000:
                break
        if content_lines:
            sections.append((line, "\n".join(content_lines).strip()))

    if not sections:
        return empty

    titles = _unique_strings(title for title, _ in sections)
    excerpt = "\n\n".join(f"{title}\n{content}" for title, content in sections)[:2_000]
    result = {**empty, "section_titles": titles, "source_excerpt": excerpt}
    result["summary"] = _summary_from_excerpt(excerpt)
    result["matching_rules"] = _build_overview_rules(result)
    return result


def normalize_overview_context(value: object, raw_text: str) -> dict:
    """Keep the AI extraction grounded in locally detected overview source text."""
    fallback = extract_overview_context(raw_text)
    source = value if isinstance(value, dict) else {}
    overview = {
        "section_titles": _unique_strings([*fallback["section_titles"], *_string_list(source.get("section_titles"))]),
        "summary": _clean_text(source.get("summary"), limit=600) or fallback["summary"],
        "core_topics": _string_list(source.get("core_topics")),
        "target_audience": _string_list(source.get("target_audience")),
        "delivery_formats": _string_list(source.get("delivery_formats")),
        "required_roles": _string_list(source.get("required_roles")),
        "scope": _string_list(source.get("scope")),
        "expected_outcomes": _string_list(source.get("expected_outcomes")),
        "preferred_experience": _string_list(source.get("preferred_experience")),
        "source_excerpt": fallback["source_excerpt"] or _clean_text(source.get("source_excerpt"), limit=2_000),
    }
    if not overview["summary"] and overview["source_excerpt"]:
        overview["summary"] = _summary_from_excerpt(overview["source_excerpt"])
    overview["matching_rules"] = _build_overview_rules(overview)
    return overview


def _string_list(value: object, limit: int = 12) -> list[str]:
    if not isinstance(value, list):
        return []
    return _unique_strings(_clean_text(item, limit=160) for item in value)[:limit]


def _unique_strings(values) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        if not value:
            continue
        text = str(value).strip()
        key = re.sub(r"\s+", "", text).casefold()
        if key and key not in seen:
            seen.add(key)
            result.append(text)
    return result


def _clean_text(value: object, limit: int) -> str:
    return str(value).strip()[:limit] if isinstance(value, str) else ""


def _summary_from_excerpt(excerpt: str) -> str:
    return re.sub(r"\s+", " ", excerpt).strip()[:400]


def _build_overview_rules(overview: dict) -> list[dict]:
    labels = {
        "core_topics": "핵심 교육 주제",
        "target_audience": "교육 대상",
        "delivery_formats": "운영 방식",
        "required_roles": "요구 역할",
        "scope": "업무 범위",
        "expected_outcomes": "기대 성과",
        "preferred_experience": "우대 경험",
    }
    return [
        {"key": key, "label": label, "values": overview.get(key, [])}
        for key, label in labels.items()
        if overview.get(key)
    ]


def _extract_text_from_hwp(file_content: bytes, file_name: str) -> str | None:
    """한글 프로그램이나 PDF 변환 없이 HWP 5.x 본문을 로컬에서 추출합니다."""
    try:
        from matching_core.parser import parse_document

        text = parse_document(file_content, file_name).raw_text.strip()
        if _is_usable_extracted_text(text):
            logger.info("hwp_local_extraction_success", text_length=len(text))
            return text
        logger.warning("hwp_local_extraction_low_quality", text_length=len(text))
    except Exception as exc:
        logger.warning("hwp_local_extraction_failed", error=str(exc))
    return None


def _is_usable_extracted_text(text: str) -> bool:
    """바이너리 노이즈를 JSON 추출 모델에 전달하지 않도록 최소 품질을 검증합니다."""
    normalized = text.strip()
    if len(normalized) < 80:
        return False

    meaningful = sum(
        character.isalnum() or character in " \n\r\t.,:;()[]{}<>/-_+%"
        for character in normalized
    )
    return meaningful / len(normalized) >= 0.75


def _extract_text_from_docx(file_content: bytes) -> str | None:
    """python-docx로 DOCX 텍스트 추출."""
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
        if text and len(text.strip()) > 50:
            return text
    except Exception as e:
        logger.warning("docx_extraction_failed", error=str(e))
    return None


async def parse_document_with_ai(file_content: bytes, file_name: str) -> dict:
    """파일을 AI로 파싱하여 신청자격/평가기준을 추출합니다.

    전략:
    - DOCX: python-docx → 텍스트 → Gemini
    - HWP: 서버리스 호환 OLE/HWP 레코드 파서 → Gemini
    - PDF: Gemini File API에 직접 업로드

    Returns:
        dict: {"qualifications": [...], "evaluation_criteria": [...], "overview": {...}, "raw_text": str}
    """
    import asyncio

    ext = file_name.rsplit('.', 1)[-1].lower() if '.' in file_name else ''
    raw_text = ""
    use_file_api = False

    # 1. 텍스트 추출 시도
    if ext == 'docx':
        raw_text = _extract_text_from_docx(file_content) or ""
    elif ext == 'hwp':
        raw_text = _extract_text_from_hwp(file_content, file_name) or ""
        if not raw_text:
            # Gemini File API는 HWP 이진 형식을 안정적으로 처리하지 못하므로
            # 로컬 추출 실패를 원인과 함께 반환한다.
            return {
                "qualifications": [],
                "evaluation_criteria": [],
                "overview": extract_overview_context(""),
                "raw_text": "",
                "parse_error": "HWP 본문을 추출하지 못했습니다. 암호화 또는 배포용 문서인지 확인해 주세요.",
            }
    elif ext == 'pdf':
        # PDF는 Gemini File API로 직접 처리
        use_file_api = True
    else:
        # 기타: matching_core 파서 시도
        try:
            from matching_core.parser import parse_document
            doc = parse_document(file_content, file_name)
            raw_text = doc.raw_text or ""
        except:
            pass

    # 2. Gemini로 구조화 추출
    try:
        client = _get_client()
        loop = asyncio.get_event_loop()

        if use_file_api:
            # PDF: File API로 직접 업로드
            result = await loop.run_in_executor(None, lambda: _parse_with_file_api(client, file_content, file_name))
        elif raw_text and len(raw_text) > 50:
            # 텍스트가 있으면 텍스트 기반 파싱
            result = await loop.run_in_executor(None, lambda: _parse_with_text(client, raw_text))
        else:
            # 텍스트 추출 실패 → File API 시도
            result = await loop.run_in_executor(None, lambda: _parse_with_file_api(client, file_content, file_name))

        result["raw_text"] = raw_text or result.get("raw_text", "")
        result["overview"] = normalize_overview_context(result.get("overview"), result["raw_text"])
        return result

    except Exception as e:
        logger.error("ai_parsing_failed", error=str(e))
        return {
            "qualifications": [],
            "evaluation_criteria": [],
            "overview": extract_overview_context(raw_text),
            "raw_text": raw_text,
        }


def _parse_with_text(client: genai.Client, text: str) -> dict:
    """텍스트를 Gemini에 전달하여 파싱."""
    overview_hint = extract_overview_context(text)
    prompt = EXTRACTION_PROMPT + "\n\n## 과업지시서 텍스트:\n" + text[:10000]
    if overview_hint["source_excerpt"] and overview_hint["source_excerpt"] not in text[:10000]:
        prompt += "\n\n## 제목 규칙으로 찾은 개요 후보 (반드시 원문과 대조):\n" + overview_hint["source_excerpt"]

    response = client.models.generate_content(model=MODEL, contents=prompt)
    return _parse_json_response(response.text)


def _parse_with_file_api(client: genai.Client, file_content: bytes, file_name: str) -> dict:
    """Gemini File API로 파일을 직접 업로드하여 파싱."""
    ext = file_name.rsplit('.', 1)[-1].lower()
    mime_map = {'pdf': 'application/pdf', 'hwp': 'application/x-hwp', 'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'}
    mime_type = mime_map.get(ext, 'application/octet-stream')

    # 임시 파일 저장
    with tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False) as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name

    try:
        # Gemini File API에 업로드
        uploaded_file = client.files.upload(file=tmp_path, config=types.UploadFileConfig(mime_type=mime_type))
        logger.info("file_uploaded_to_gemini", name=uploaded_file.name)

        # AI로 파싱 요청
        response = client.models.generate_content(
            model=MODEL,
            contents=[
                types.Content(parts=[
                    types.Part.from_uri(file_uri=uploaded_file.uri, mime_type=mime_type),
                    types.Part.from_text(text=EXTRACTION_PROMPT),
                ])
            ]
        )

        result = _parse_json_response(response.text)

        # 원문 텍스트도 추출 시도
        try:
            text_response = client.models.generate_content(
                model=MODEL,
                contents=[
                    types.Content(parts=[
                        types.Part.from_uri(file_uri=uploaded_file.uri, mime_type=mime_type),
                        types.Part.from_text(text="이 문서의 전체 텍스트를 그대로 출력해주세요. 다른 설명 없이 문서 내용만:"),
                    ])
                ]
            )
            result["raw_text"] = text_response.text[:10000]
        except:
            pass

        return result
    finally:
        os.unlink(tmp_path)


def _parse_json_response(text: str) -> dict:
    """AI 응답에서 JSON을 추출."""
    text = text.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:])
        if "```" in text:
            text = text[:text.rfind("```")]
    text = text.strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # JSON 부분만 추출 시도
        start = text.find("{")
        end = text.rfind("}") + 1
        if start >= 0 and end > start:
            try:
                return json.loads(text[start:end])
            except:
                pass
    return {"qualifications": [], "evaluation_criteria": [], "overview": {}}


async def generate_match_reason(
    task_description: str,
    instructor_name: str,
    instructor_keywords: list[str],
    instructor_experience: int,
    score: float,
    keyword_score: float,
    qual_score: float,
    exp_score: float,
) -> str:
    """AI로 매칭 추천 이유를 3문장 글머리 기호로 생성합니다."""
    import asyncio

    prompt = f"""당신은 교육 강사 매칭 전문가입니다.
아래 정보를 바탕으로 이 강사가 해당 과업에 적합한 이유를 분석해주세요.

## 과업 요약:
{task_description[:500]}

## 강사 정보:
- 이름: {instructor_name}
- 보유기술: {', '.join(instructor_keywords[:15])}
- 강의건수: {instructor_experience}건

## 매칭 점수:
- 종합: {score}점/100
- 키워드 일치: {keyword_score}/40
- 자격 매칭: {qual_score}/30
- 경력 매칭: {exp_score}/30

## 작성 규칙:
- 한국어로 작성
- 아래 JSON 형식으로만 출력:

{{"strengths": ["강점 1문장", "강점 1문장"], "weaknesses": ["약점 1문장"], "summary": "종합 평가 1문장"}}

- strengths: 이 강사가 해당 과업에 적합한 구체적 이유 2가지 (각 1문장)
- weaknesses: 부족한 점 1가지 (1문장)
- summary: 종합 평가 1문장

JSON만 출력하세요:"""

    try:
        client = _get_client()
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(
            None,
            lambda: client.models.generate_content(model=MODEL, contents=prompt)
        )
        text = response.text.strip()
        # JSON 파싱 시도
        parsed = _parse_json_response(text)
        if parsed and ("strengths" in parsed or "weaknesses" in parsed):
            return json.dumps(parsed, ensure_ascii=False)
        return text
    except Exception as e:
        logger.error("ai_reason_failed", error=str(e))
        return json.dumps({"strengths": [], "weaknesses": [], "summary": "AI 분석을 수행할 수 없습니다."}, ensure_ascii=False)


async def ai_score_instructors(
    task_requirements: dict,
    instructors: list[dict],
) -> list[dict]:
    """AI로 강사별 매칭 점수를 차별화하여 부여합니다.

    Args:
        task_requirements: {"qualifications": [...], "evaluation_criteria": [...]}
        instructors: [{"id": str, "name": str, "keywords": [...], "experience_years": int, ...}]

    Returns:
        [{"instructor_id": str, "score": int, "reason": str}] (점수 내림차순)
    """
    import asyncio

    # 과업 요약
    task_summary_parts = []
    for q in task_requirements.get("qualifications", [])[:5]:
        task_summary_parts.append(f"자격: {q.get('description', '')}")
    for e in task_requirements.get("evaluation_criteria", [])[:5]:
        task_summary_parts.append(f"평가: {e.get('description', '')}")
    task_summary = "\n".join(task_summary_parts)

    # 강사 목록 요약 (최대 30명씩 배치)
    instructor_summaries = []
    for i in instructors[:30]:
        keywords = ", ".join((i.get("keywords") or [])[:10])
        instructor_summaries.append(f"- ID:{i['id'][:8]} | {i['name']} | 기술:{keywords} | 강의:{i.get('experience_years',0)}건")

    instructors_text = "\n".join(instructor_summaries)

    prompt = f"""당신은 교육 강사 매칭 AI입니다.
아래 과업지시서의 요구사항과 강사 목록을 비교하여, 각 강사에게 0~100점 매칭 점수를 부여해주세요.

## 과업 요구사항:
{task_summary}

## 강사 목록:
{instructors_text}

## 점수 기준:
- 과업에서 요구하는 기술/분야와 강사의 보유기술이 얼마나 일치하는가
- 강의 경험(건수)이 풍부한가
- 관련 분야 전문성이 있는가
- 점수는 0~100 사이, 강사별로 반드시 다르게 (동점 최소화)

## 출력 형식 (JSON 배열만 출력):
[{{"id":"첫8자리","score":85}},{{"id":"첫8자리","score":72}}]

JSON만 출력:"""

    try:
        client = _get_client()
        loop = asyncio.get_event_loop()
        response = await loop.run_in_executor(
            None,
            lambda: client.models.generate_content(model=MODEL, contents=prompt)
        )

        result = _parse_json_response(response.text)
        if isinstance(result, list):
            return result
        elif isinstance(result, dict) and "scores" in result:
            return result["scores"]
        return []
    except Exception as e:
        logger.error("ai_scoring_failed", error=str(e))
        return []
